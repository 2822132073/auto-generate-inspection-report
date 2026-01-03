"""
报告生成服务
生成 DOCX 格式的巡检报告
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path
from config import REPORTS_DIR, SCREENSHOTS_DIR, REPORT_TITLE_DEFAULT, REPORT_SCREENSHOT_WIDTH
from models.database import get_db_connection
from services.inspection_service import InspectionService
from services.template_service import TemplateService
import json


class ReportService:
    """报告生成服务"""

    @staticmethod
    def generate_project_report(project_id, options=None, template_id=None):
        """
        生成项目巡检报告

        Args:
            project_id: 项目ID
            options: 可选配置（title, include_screenshots）
            template_id: 模板ID（可选，不指定则使用默认模板）

        Returns:
            dict: 报告信息（路径、大小等）
        """
        options = options or {}

        # 获取模板配置
        template_config = None
        if template_id:
            template = TemplateService.get_template(template_id)
            if not template:
                raise ValueError(f"模板 {template_id} 不存在")
            template_config = template['config']
        else:
            # 使用默认模板
            default_template = TemplateService.get_default_template()
            if default_template:
                template_config = default_template['config']
        
        # 如果没有模板配置，使用传统方式（向后兼容）
        if not template_config:
            template_config = {
                'include_return_code': True,
                'include_output': True,
                'include_screenshots': options.get('include_screenshots', True),
                'title_format': '{title}',
                'section_organization': 'by_host'
            }

        # 获取项目下所有主机的最新巡检记录
        host_records = InspectionService.get_project_latest_hosts(project_id)

        if not host_records:
            raise ValueError(f"项目 {project_id} 没有巡检记录")

        # 创建文档
        doc = Document()

        # 设置中文字体支持
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        style.font.size = Pt(10.5)

        # 添加标题（使用模板格式）
        title_text = template_config['title_format'].format(
            project_id=project_id,
            title=options.get('title', REPORT_TITLE_DEFAULT)
        )
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加项目信息
        doc.add_paragraph(f"项目ID: {project_id}")
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"主机数量: {len(host_records)}")
        doc.add_paragraph()  # 空行

        # 按主机分章节
        include_screenshots = template_config.get('include_screenshots', True)
        include_return_code = template_config.get('include_return_code', True)
        include_output = template_config.get('include_output', True)
        show_command_title = template_config.get('show_command_title', True)

        for host_idx, record in enumerate(host_records, 1):
            # 获取主机详细信息（包含命令）
            detail = InspectionService.get_inspection_detail(record['id'])

            if not detail:
                continue

            # 主机章节标题
            doc.add_heading(f"{ReportService._to_chinese_number(host_idx)}、主机 ({detail['hostname']})", level=1)

            # 1.1 基本信息
            doc.add_heading(f"{host_idx}.1 基本信息", level=2)

            info_table = doc.add_table(rows=6, cols=2)
            info_table.style = 'Light Grid Accent 1'

            info_data = [
                ('主机名', detail.get('hostname', '-')),
                ('IP地址', detail.get('ip', '-')),
                ('操作系统', detail.get('os', '-')),
                ('内核版本', detail.get('kernel', '-')),
                ('架构', detail.get('arch', '-')),
                ('巡检时间', detail.get('timestamp', '-'))
            ]

            for row_idx, (label, value) in enumerate(info_data):
                cells = info_table.rows[row_idx].cells
                cells[0].text = label
                cells[1].text = str(value)

            # 1.2 系统信息汇总
            doc.add_heading(f"{host_idx}.2 系统信息汇总", level=2)

            commands = detail.get('commands', [])

            for cmd_idx, cmd in enumerate(commands, 1):
                # 命令小标题（根据模板配置决定是否显示）
                if show_command_title:
                    doc.add_heading(f"{host_idx}.2.{cmd_idx} {cmd['command']}", level=3)

                # 返回码（根据模板配置）
                if include_return_code:
                    return_code = cmd.get('return_code', 0)
                    return_code_para = doc.add_paragraph(f"返回码: {return_code}")
                    if return_code != 0:
                        return_code_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)

                # 输出内容（根据模板配置）
                if include_output:
                    output = cmd.get('output', '')
                    if output:
                        doc.add_paragraph('执行结果:')
                        output_para = doc.add_paragraph(output)
                        output_para.style = 'No Spacing'
                        run = output_para.runs[0]
                        run.font.name = 'Consolas'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                        run.font.size = Pt(9)

                # 插入截图
                if include_screenshots and cmd.get('screenshot_path'):
                    screenshot_file = SCREENSHOTS_DIR / cmd['screenshot_path']
                    if screenshot_file.exists():
                        doc.add_paragraph('终端截图:')
                        try:
                            doc.add_picture(str(screenshot_file), width=Inches(REPORT_SCREENSHOT_WIDTH))
                        except Exception as e:
                            doc.add_paragraph(f'[截图加载失败: {e}]')

                doc.add_paragraph()  # 空行

        # 页脚：生成时间
        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].text = f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存文档
        month_dir = datetime.now().strftime('%Y-%m')
        save_dir = REPORTS_DIR / month_dir
        save_dir.mkdir(exist_ok=True, parents=True)

        filename = f"project_{project_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        file_path = save_dir / filename

        doc.save(str(file_path))

        # 记录到数据库
        file_size = file_path.stat().st_size
        relative_path = f"{month_dir}/{filename}"

        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO report_generations (project_id, report_path, format, file_size, host_count)
                VALUES (?, ?, 'docx', ?, ?)
            ''', (project_id, relative_path, file_size, len(host_records)))
            report_id = cursor.lastrowid

        return {
            'report_id': report_id,
            'file_path': relative_path,
            'file_size': file_size,
            'host_count': len(host_records),
            'generated_at': datetime.now().isoformat()
        }

    @staticmethod
    def _to_chinese_number(num):
        """将数字转换为中文数字（一、二、三...）"""
        chinese_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if num <= 10:
            return chinese_nums[num]
        elif num < 20:
            return f"十{chinese_nums[num - 10]}"
        else:
            tens = num // 10
            ones = num % 10
            if ones == 0:
                return f"{chinese_nums[tens]}十"
            else:
                return f"{chinese_nums[tens]}十{chinese_nums[ones]}"

    @staticmethod
    def get_latest_project_report(project_id):
        """获取项目的最新报告路径"""
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT report_path FROM report_generations
                WHERE project_id = ?
                ORDER BY generated_at DESC
                LIMIT 1
            ''', (project_id,))

            row = cursor.fetchone()
            return row['report_path'] if row else None
